"""
Extractor de DOCX usando python-docx.

Extrae:
- Texto por párrafo preservando estructura
- Encabezados con jerarquía (Heading 1, 2, 3)
- Tablas con estructura
- Metadata del documento
"""

import logging
from typing import List, Dict, Optional
from pathlib import Path
from docx import Document

_logger = logging.getLogger(__name__)


class DOCXExtractor:
    """
    Extractor especializado para archivos DOCX.
    
    Usa python-docx para:
    - Extracción de texto preservando estructura
    - Detección de encabezados y jerarquía
    - Extracción de tablas
    - Metadata del documento
    """
    
    def __init__(self):
        """Inicializa el extractor DOCX."""
        _logger.info("📄 DOCXExtractor inicializado")
    
    def extract(self, docx_path: str) -> Dict:
        """
        Extrae todo el contenido de un DOCX.
        
        Args:
            docx_path: Ruta local del archivo DOCX
        
        Returns:
            Diccionario con:
            {
                "metadata": {...},
                "sections": [
                    {
                        "type": "heading|paragraph|table",
                        "level": 1,  # solo para headings
                        "text": "...",
                        "table": {...}  # solo para tables
                    },
                    ...
                ],
                "extraction_method": "python-docx"
            }
        
        Raises:
            ValueError: Si el archivo no es válido
            FileNotFoundError: Si el archivo no existe
        """
        try:
            _logger.info(f"📖 Extrayendo DOCX: {docx_path}")
            
            if not Path(docx_path).exists():
                raise FileNotFoundError(f"Archivo no encontrado: {docx_path}")
            
            doc = Document(docx_path)
            
            # Extraer metadata
            metadata = self._extract_metadata(doc)
            
            # Extraer contenido por secciones
            sections = []
            
            # Procesar párrafos y encabezados
            for para in doc.paragraphs:
                section = self._process_paragraph(para)
                if section:
                    sections.append(section)
            
            # Procesar tablas
            for i, table in enumerate(doc.tables):
                table_data = self._extract_table(table, i)
                sections.append({
                    "type": "table",
                    "table_index": i,
                    "table": table_data
                })
            
            result = {
                "metadata": metadata,
                "sections": sections,
                "total_paragraphs": len(doc.paragraphs),
                "total_tables": len(doc.tables),
                "extraction_method": "python-docx"
            }
            
            _logger.info(
                f"✅ DOCX extraído: {len(sections)} secciones, "
                f"{len(doc.tables)} tablas"
            )
            
            return result
        
        except Exception as e:
            _logger.error(f"❌ Error extrayendo DOCX: {e}", exc_info=True)
            raise
    
    def _extract_metadata(self, doc: Document) -> Dict:
        """
        Extrae metadata del documento.
        
        Args:
            doc: Objeto Document
        
        Returns:
            Metadata del documento
        """
        try:
            core_props = doc.core_properties
            return {
                "title": core_props.title or "",
                "author": core_props.author or "",
                "subject": core_props.subject or "",
                "created": str(core_props.created) if core_props.created else "",
                "modified": str(core_props.modified) if core_props.modified else ""
            }
        except Exception as e:
            _logger.warning(f"⚠️ Error extrayendo metadata: {e}")
            return {}
    
    def _process_paragraph(self, para) -> Optional[Dict]:
        """
        Procesa un párrafo.
        
        Args:
            para: Objeto Paragraph de python-docx
        
        Returns:
            Sección procesada o None si está vacío
        """
        text = para.text.strip()
        
        if not text:
            return None
        
        # Detectar si es encabezado
        if para.style.name.startswith('Heading'):
            try:
                level = int(para.style.name.split()[-1])
            except (ValueError, IndexError):
                level = 1
            
            return {
                "type": "heading",
                "level": level,
                "text": text
            }
        
        # Es un párrafo normal
        return {
            "type": "paragraph",
            "text": text
        }
    
    def _extract_table(self, table, table_index: int) -> Dict:
        """
        Extrae una tabla.
        
        Args:
            table: Objeto Table de python-docx
            table_index: Índice de la tabla
        
        Returns:
            {
                "headers": [...],
                "rows": [[...]],
                "num_rows": 10,
                "num_cols": 5
            }
        """
        try:
            # Primera fila como headers
            headers = []
            if len(table.rows) > 0:
                headers = [cell.text.strip() for cell in table.rows[0].cells]
            
            # Resto como datos
            rows = []
            for row in table.rows[1:]:
                row_data = [cell.text.strip() for cell in row.cells]
                rows.append(row_data)
            
            return {
                "headers": headers,
                "rows": rows,
                "num_rows": len(rows),
                "num_cols": len(headers)
            }
        
        except Exception as e:
            _logger.error(f"❌ Error extrayendo tabla {table_index}: {e}")
            return {
                "headers": [],
                "rows": [],
                "num_rows": 0,
                "num_cols": 0,
                "error": str(e)
            }
    
    def extract_text(self, docx_path: str) -> str:
        """
        Extrae solo el texto plano (sin estructura).
        
        Args:
            docx_path: Ruta del DOCX
        
        Returns:
            Texto completo del documento
        """
        full_extraction = self.extract(docx_path)
        
        text_parts = []
        for section in full_extraction["sections"]:
            if section["type"] in ["heading", "paragraph"]:
                text_parts.append(section["text"])
        
        return "\n\n".join(text_parts)
    
    def extract_headings(self, docx_path: str) -> List[Dict]:
        """
        Extrae solo los encabezados con jerarquía.
        
        Args:
            docx_path: Ruta del DOCX
        
        Returns:
            Lista de encabezados: [{"level": 1, "text": "..."}]
        """
        full_extraction = self.extract(docx_path)
        
        return [
            {"level": s["level"], "text": s["text"]}
            for s in full_extraction["sections"]
            if s["type"] == "heading"
        ]
    
    def extract_all_text(self, docx_path: str) -> str:
        """
        Extrae todo el texto del documento (compatible con XLSXExtractor).
        
        Args:
            docx_path: Ruta del DOCX
        
        Returns:
            Texto completo del documento
        """
        return self.extract_text(docx_path)


if __name__ == "__main__":
    # Test básico
    import sys
    logging.basicConfig(level=logging.INFO)
    
    if len(sys.argv) > 1:
        extractor = DOCXExtractor()
        result = extractor.extract(sys.argv[1])
        print(f"\n✅ Extraído: {result['total_paragraphs']} párrafos, {result['total_tables']} tablas")
        
        # Mostrar primeros 3 elementos
        for section in result['sections'][:3]:
            if section['type'] == 'heading':
                print(f"\n{'#' * section['level']} {section['text']}")
            elif section['type'] == 'paragraph':
                print(f"\n{section['text'][:100]}...")
            elif section['type'] == 'table':
                print(f"\n[Tabla {section['table_index']}: {section['table']['num_rows']}x{section['table']['num_cols']}]")
    else:
        print("Uso: python docx_extractor.py <ruta_docx>")

